from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".markdown", ".docx"}


class UnsupportedFileTypeError(ValueError):
    pass


def is_supported_file(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_EXTENSIONS


def load_document(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    if suffix in {".txt", ".md", ".markdown"}:
        return _load_text(path)
    raise UnsupportedFileTypeError(
        f"Unsupported file type: {path.suffix}. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def _load_text(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8").strip()
    if not text:
        return []
    return [Document(page_content=text, metadata={"source": path.name, "path": str(path)})]


def _load_pdf(path: Path) -> list[Document]:
    reader = PdfReader(str(path))
    documents: list[Document] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": path.name,
                        "path": str(path),
                        "page": page_index,
                    },
                )
            )
    return documents


def _load_docx(path: Path) -> list[Document]:
    doc = DocxDocument(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs]
    text = "\n\n".join(paragraph for paragraph in paragraphs if paragraph)
    if not text:
        return []
    return [Document(page_content=text, metadata={"source": path.name, "path": str(path)})]
